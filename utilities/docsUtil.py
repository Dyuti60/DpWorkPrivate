import docx
import os
from docx.shared import Inches

def createDocxFile():
    document = docx.Document()
    return document

def addMainHeading(document, documentTitle):
    document.add_heading(str(documentTitle),0)

def addMediumHeading(document, documentTitle):
    document.add_heading(str(documentTitle),2)

def addSmallHeading(document, documentTitle):
    document.add_heading(str(documentTitle),3)

def appendContent(document, content):
    document.add_paragraph(str(content))

def insertImageInDocx(document, image):
    p=document.add_paragraph()
    r=p.add_run()
    r.add_picture(image, width=Inches(6.5), height=Inches(3.5))

def saveDocument(document, saveFilename):
    os.chdir("../")
    desiredLocation=os.getcwd()+'\\documentation\\dpWorkLoginPageDocumentation\\{}'.format(str(saveFilename))
    document.save("{}.docx".format(str(desiredLocation)))

